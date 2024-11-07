import os
import json
import google.generativeai as genai
import openpyxl
import logging
import argparse
from openpyxl.styles import Alignment
from openpyxl.utils.datetime import from_excel
from dotenv import load_dotenv
import subprocess
import concurrent.futures
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import sys
from pathlib import Path
import time
import google.api_core.exceptions
from docx import Document
import datetime
import xml.parsers.expat
import webbrowser
from dateutil import parser

# ユーザーディレクトリのDocumentsフォルダのパスを取得
documents_path = Path.home() / "Documents"
log_file_path = documents_path / "app_log.txt"

# ログの設定
logging.basicConfig(
    level=logging.DEBUG,  # ログレベルをDEBUGに設定
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file_path),  # ログファイルのパスを指定
        logging.StreamHandler()  # コンソールにも出力
    ]
)

def get_current_dir():
    # この関数は、現在のスクリプトがどこにあるかを教えてくれます。
    if getattr(sys, 'frozen', False):
        # もしプログラムがPyInstallerでパッケージ化されているなら
        return Path(sys._MEIPASS)  # 特別なフォルダを使います
    else:
        # そうでないなら、開発中のフォルダを使います
        return Path(__file__).resolve().parent

settings_path = os.path.join(get_current_dir(), 'settings.json')
print(f"Settings path: {settings_path}")
print(f"File exists: {os.path.exists(settings_path)}")

# 現在のスクリプトのディレクトリを取得
current_dir = Path(__file__).resolve().parent

# 環境変数の読み込み
load_dotenv(current_dir / '環境変数.env')

# プロジェクトディレクトリの設定
project_dir = os.path.dirname(os.path.abspath(__file__))

# APIキーの設定
API_KEYS = [os.getenv(f'GEMINI_API_KEY_{i}') for i in range(1, 11)]  # 10個のAPIキーを取得

# 処理済みファイルのログファイル
PROCESSED_FILES_LOG = os.path.join(current_dir, 'processed_files.json')

def load_processed_files():
    # この関数は、すでに処理したファイルのリストを読み込みます。
    if os.path.exists(PROCESSED_FILES_LOG):
        # もしログファイルが存在するなら
        with open(PROCESSED_FILES_LOG, 'r') as f:
            return json.load(f)  # ファイルを開いて内容を読み込みます
    return {}  # ファイルがなければ空のリストを返します

def save_processed_files(processed_files):
    # この関数は、処理したファイルのリストを保存します。
    with open(PROCESSED_FILES_LOG, 'w') as f:
        json.dump(processed_files, f, indent=2)  # ファイルに書き込みます

def get_unprocessed_audio_files():
    # この関数は、まだ処理していない音声ファイルを探します。
    processed_files = load_processed_files()  # すでに処理したファイルを取得
    audio_files = [f for f in os.listdir(current_dir) if f.endswith('.mp3')]
    # フォルダ内のすべての.mp3ファイルをリストにします
    return [f for f in audio_files if f not in processed_files]
    # まだ処理していないファイルだけを返します

def create_extraction_prompt(text):
    # この関数は、会議の内容から情報を抽出するための指示を作ります。
    return f"""
    この文章はとある会議の内容です。
    以下の文章から、次の項目を抽出してください：
    1. 議題①
    2. 議題①の要約
    3. 議題②
    4. 議題②の要約
    5. 議題③
    6. 議題③の要約
    7. 議題④
    8. 議題④の要約
    9. 議題⑤
    10. 議題⑤の要約
    11. 議題⑥
    12. 議題⑥の要約
    13. 議題⑦
    14. 議題⑦の要約
    15. 議題⑧
    16. 議題⑧の要約
    17. 議題⑨
    18. 議題⑨の要約
    19. 議題⑩
    20. 議題⑩の要約

    抽出する際は、必ず以下の形式で出力してください：
    議題①: [議題の内容]
    議題①の要約: [要約内容]

    議題②: [議題の内容]
    議題②の要約: [要約内容]

    議題③: [議題の内容]
    議題③の要約: [要約内容]

    ...

    議題⑩: [議題の内容]
    議題⑩の要約: [要約内容]

    注意事項:
    - 各議題とその要約を必ず上記の形式で出力してください。
    - 議題が10個未満の場合は、存在する議題のみを抽出してください。
    - 要約は簡潔かつ具体的にしてください。
    - 議題の番号（①、②など）は必ず付けてください。
    - 各行は必ず「議題○:」または「議題○の要約:」で始まるようにしてください。
    - 議題や要約の前に「*」や「**」などの記号を付けないでください。
    - 議題というのはあくまで表現の一つであり、会話内容が議事録形式で記されていれば構いません。インタビューの文章等からも適切に議題を抽出してください。
    - インタビューのような文章であっても、適切に議題を抽出してください。

    文章:
    {text}
    """

def get_ffmpeg_path():
    """ffmpegのパスを取得する関数"""
    # この関数は、ffmpegというプログラムがどこにあるかを教えてくれます

    if hasattr(sys, '_MEIPASS'):
        # もしプログラムがPyInstallerでパッケージ化されているなら
        return Path(sys._MEIPASS) / "ffmpeg"  # 特別なフォルダの中にあるffmpegを使います
    else:
        # そうでないなら（普通に開発しているとき）
        return Path(__file__).resolve().parent / "ffmpeg"  # このプログラムと同じフォルダにあるffmpegを使います

def get_ffprobe_path():
    """ffprobeのパスを取得する関数"""
    if hasattr(sys, '_MEIPASS'):
        return Path(sys._MEIPASS) / "ffprobe"  # PyInstallerでパッケージ化された場合のffprobeのパス
    else:
        return Path(__file__).resolve().parent / "ffprobe"  # 開発環境でのffprobeのパス

def split_audio_file(audio_file_path, num_parts):
    """音声ファイルを指定された数の部分に重なりを持たせて分割する関数"""
    # この関数は、長い音声ファイルを小さな部分に分けます。
    # 分けた部分は少し重なりを持つので、途切れないようになっています。

    file_size = os.path.getsize(audio_file_path)  # ファイルの大きさを調べます
    duration = get_audio_duration(audio_file_path)  # 音声ファイルの長さを取得します
    part_duration = duration / num_parts  # 各部分の長さを計算します
    overlap_duration = part_duration * 0.1  # 10%の重なりを持たせます

    parts = []  # 分割した音声ファイルのリストを作ります
    for i in range(num_parts):
        # 各部分の開始時間を計算します
        start_time = max(0, i * part_duration - (overlap_duration if i > 0 else 0))
        # 新しい音声ファイルの名前を決めます
        part_file = f"{audio_file_path}_part{i+1}.mp3"  # 拡張子をmp3のままにします

        # 音声ファイルの種類に応じて、分割の方法を変えます
        if audio_file_path.endswith('.mp3'):
            # MP3ファイルの場合の分割方法
            command = [
                str(get_ffmpeg_path()),  # ffmpegというソフトウェアのパスを取得します
                '-y',  # 同じ名前のファイルがあれば上書きします
                '-i', audio_file_path,  # 元の音声ファイルを指定します
                '-ss', str(start_time),  # 開始時間を指定します
                '-t', str(part_duration + (overlap_duration if i < num_parts - 1 else 0)),  # 部分の長さを指定します
                '-c', 'copy',  # 音声をそのままコピーします（音質を変えません）
                part_file  # 新しい音声ファイルの名前を指定します
            ]
        elif audio_file_path.endswith('.m4a'):
            # M4Aファイルの場合の分割方法
            part_file = f"{audio_file_path}_part{i+1}.m4a"  # 拡張子をm4aのままにします
            command = [
                str(get_ffmpeg_path()),
                '-y',
                '-i', audio_file_path,
                '-ss', str(start_time),
                '-t', str(part_duration + (overlap_duration if i < num_parts - 1 else 0)),
                '-c', 'copy',
                part_file
            ]
        elif audio_file_path.endswith('.wav'):
            # WAVファイルの場合の分割方法
            part_file = f"{audio_file_path}_part{i+1}.wav"  # 拡張子をwavのままにします
            command = [
                str(get_ffmpeg_path()),
                '-y',
                '-i', audio_file_path,
                '-ss', str(start_time),
                '-t', str(part_duration + (overlap_duration if i < num_parts - 1 else 0)),
                '-c', 'pcm_s16le',  # WAV用の音声形式を指定します
                part_file
            ]

        # 音声ファイルを実際に分割します
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode != 0:
            # エラーが起きた場合は記録します
            logging.error(f"FFmpegエラー: {result.stderr}")
        parts.append(part_file)  # 分割したファイルをリストに追加します

    return parts  # 分割したファイルのリストを返します

def get_audio_duration(audio_file_path):
    """音声ファイルの長さを取得する関数"""
    # この関数は、音声ファイルの再生時間（長さ）を秒単位で取得します

    # ffprobeというツールを使うためのコマンドを準備します
    command = [
        str(get_ffprobe_path()),  # ffprobeのパスを取得します
        '-v', 'error',  # エラー以外の出力を抑制します
        '-show_entries', 'format=duration',  # ファイルの長さ（duration）を取得するよう指示します
        '-of', 'default=noprint_wrappers=1:nokey=1',  # 結果をシンプルな形式で出力するよう指示します
        audio_file_path  # 調べたい音声ファイルのパスを指定します
    ]

    # 準備したコマンドを実行し、結果を取得します
    result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

    # 結果から音声ファイルの長さ（秒）を取り出し、小数点の数値として返します
    return float(result.stdout.strip())

# グローバル変数の定義
transcription_prompt = ""

def load_prompt_from_settings():
    """settings.jsonからプロンプトを読み込む関数"""
    settings_path = os.path.join(get_current_dir(), 'settings.json')
    logging.info(f"Settings path: {settings_path}")  # 追加: パスをログに出力
    if os.path.exists(settings_path):
        logging.info("settings.jsonが見つかりました。")  # 追加: ファイル存在確認
        with open(settings_path, 'r', encoding='utf-8') as f:
            try:
                settings = json.load(f)
                logging.info("settings.jsonを正常に読み込みました。")  # 追加: 読み込み成功
                return settings.get('transcription_prompt', '')
            except json.JSONDecodeError as e:
                logging.error(f"JSONデコードエラー: {str(e)}")  # 追加: JSONデコードエラー
    else:
        logging.error("settings.jsonが見つかりません。")  # 追加: ファイルが見つからない場合
    return ''

def transcribe_audio_with_key(audio_file, api_key, retries=3):
    """指定されたAPIキーを使用して音声ファイルを文字起こしする関数"""
    # この関数は、音声ファイルをテキストに変換します

    # プロンプトをログに出力（1回だけ）
    if transcription_prompt:
        # もし文字起こしの指示（プロンプト）があれば、それをログに記録します
        logging.info(f"今回は以下のプロンプトで文字起こしをします:\n{transcription_prompt}")
    else:
        # プロンプトがない場合はエラーを記録して、関数を終了します
        logging.error("プロンプトが取得できませんでした。")
        return None

    # 指定された回数（デフォルトは3回）まで文字起こしを試みます
    for attempt in range(retries):
        try:
            # 音声ファイルを開いてデータを読み込みます
            with open(audio_file, 'rb') as audio:
                audio_data = audio.read()

            # Geminiモデルを設定します
            model = genai.GenerativeModel('gemini-1.5-pro')
            genai.configure(api_key=api_key)

            # モデルを使って音声データを文字に起こします
            response = model.generate_content(
                [
                    transcription_prompt,
                    {"mime_type": "audio/mp3", "data": audio_data}
                ]
            )

            # 文字起こしが成功したかチェックします
            if hasattr(response, 'text'):
                # 成功した場合、ログに記録して結果を返します
                logging.info(f"{audio_file}の文字起こしが成功しました。")
                return response.text
            else:
                # テキストが含まれていない場合はエラーを記録します
                logging.error(f"文字起こし失敗: {audio_file} - レスポンスにテキストが含まれていません。")
        except google.api_core.exceptions.ResourceExhausted:
            # APIの利用制限に達した場合のエラーを記録します
            logging.error(f"文字起こし失敗: {audio_file} - 429 Resource has been exhausted (e.g. check quota).")
        except Exception as e:
            # その他のエラーが発生した場合、エラー内容を記録します
            logging.error(f"文字起こし失敗: {audio_file} - {str(e)}")
        
        # リトライが可能な場合は、次の試行を行います
        if attempt < retries - 1:
            logging.info(f"リトライを試みます ({attempt + 2}/{retries})")
            time.sleep(60)  # 1分待ってから次の試行を行います
        else:
            # すべての試行が失敗した場合、最終的なエラーを記録します
            logging.error(f"{audio_file}の文字起こしが{retries}回失敗しました。")
    
    # すべての試行が失敗した場合はNoneを返します
    return None

def extract_information(text, api_key):
    # この関数は、テキストから重要な情報を抽出します

    # テキストの空白を整理します
    cleaned_text = " ".join(text.split())

    # APIキーがない場合はエラーを記録して終了します
    if not api_key:
        logging.error("情報抽出に使用するAPIキーが設定されていません。")
        return

    # APIキーを設定して、AIモデルを準備します
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-pro')
    
    # 情報抽出のための指示文を作ります
    prompt = create_extraction_prompt(cleaned_text)

    try:
        # 情報抽出を開始します
        logging.info("情報抽出を開始します。")
        # AIモデルに指示を送り、結果を受け取ります
        response = model.generate_content(prompt)
        # 結果のテキストから余分な空白を取り除きます
        extracted_text = response.text.strip()
        # 抽出結果を記録します
        logging.info(f"抽出結果全体: {extracted_text}")
        # 抽出したテキストを返します
        return extracted_text
    except Exception as e:
        # エラーが起きた場合、詳細を記録して再度エラーを発生させます
        logging.exception(f"情報抽出中にエラーが発生しました: {str(e)}")
        raise

def create_excel(extracted_info, output_file):
    # 新しいExcelワークブックを作成します
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "議事録"

    # 列の幅を設定します
    ws.column_dimensions['A'].width = 20  # A列（議題）の幅を20に設定
    ws.column_dimensions['B'].width = 80  # B列（内容）の幅を80に設定

    # 会議詳細情報を追加します
    meeting_details = [
        "会議名",
        "日時",
        "場所",
        "参加者",
        "欠席者"
    ]

    # 会議詳細情報をExcelに書き込みます
    for i, detail in enumerate(meeting_details, start=1):
        ws.cell(row=i, column=1, value=detail)  # A列に項目名を書き込み
        ws.cell(row=i, column=1).font = openpyxl.styles.Font(bold=True)  # 太字に設定
        ws.cell(row=i, column=1).fill = openpyxl.styles.PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")  # 背景色を設定

    row = 6  # 会議詳細情報の後から議題の書き込みを開始します

    # 抽出された情報を行ごとに分割します
    lines = extracted_info.split('\n')
    current_topic = ""
    current_summary = ""

    # 各行を処理して議題と要約を抽出します
    for line in lines:
        line = line.strip()
        if line.startswith("議題"):
            if current_topic and current_summary:
                # 前の議題を書き込みます
                ws.cell(row=row, column=1, value=current_topic)
                cell = ws.cell(row=row, column=2, value=current_summary)
                cell.alignment = Alignment(wrap_text=True)  # テキストを折り返して表示
                row += 1
            parts = line.split(':', 1)
            if len(parts) == 2:
                current_topic = parts[0].strip()
                current_summary = parts[1].strip()
            else:
                current_topic = line
                current_summary = ""
        elif "の要約" in line:
            if current_topic and "の要約:" in line:
                current_summary = line.split("の要約:", 1)[1].strip()
        elif current_summary:
            current_summary += " " + line.strip()

    # 最後の議題を書き込みます
    if current_topic and current_summary:
        ws.cell(row=row, column=1, value=current_topic)
        cell = ws.cell(row=row, column=2, value=current_summary)
        cell.alignment = Alignment(wrap_text=True)

    # セルのスタイルを設定します
    for row in ws['A1:B'+str(ws.max_row)]:
        for cell in row:
            # すべてのセルに細い枠線を設定
            cell.border = openpyxl.styles.Border(left=openpyxl.styles.Side(style='thin'), 
                                                 right=openpyxl.styles.Side(style='thin'), 
                                                 top=openpyxl.styles.Side(style='thin'), 
                                                 bottom=openpyxl.styles.Side(style='thin'))
            if cell.column == 1:  # A列（議題）のセルの場合
                cell.font = openpyxl.styles.Font(bold=True)  # 太字に設定
                cell.fill = openpyxl.styles.PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")  # 背景色を設定
            elif cell.column == 2:  # B列（内容）のセルの場合
                cell.alignment = Alignment(wrap_text=True)  # テキストを折り返して表示

    # B列の幅を内容に合わせて自動調します
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) for cell in column_cells)
        if column_cells[0].column_letter == 'B':
            ws.column_dimensions[column_cells[0].column_letter].width = min(100, max(80, length))

    # Excelファイルを保存します
    try:
        wb.save(output_file)
        logging.info(f"Excelファイルが正常に作成されました: {output_file}")
    except PermissionError:
        logging.error(f"Excelファイルの保存に失敗しました。書き込み権限がありません: {output_file}")
    except Exception as e:
        logging.error(f"Excelファイルの保存中にエラーが発生しました: {str(e)}")

def load_output_directory():
    """settings.jsonから出力先ディレクトリを読み込む関数"""
    settings_path = os.path.join(get_current_dir(), 'settings.json')
    if os.path.exists(settings_path):
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
            return settings.get('output_directory', os.path.join(Path.home(), 'Documents'))
    return os.path.join(Path.home(), 'Documents')

def process_audio_file(audio_file_path, processed_files):
    try:
        audio_file_name = os.path.basename(audio_file_path)
        file_size = os.path.getsize(audio_file_path)
        logging.info(f"{audio_file_name}の処理を開始します。ファイルサイズ: {file_size / (1024 * 1024):.2f}MB")

        # APIキーをロード
        api_keys = load_api_keys()
        if not api_keys:
            logging.error("APIキーがロードされていません。処理を中止します。")
            return False

        # 音声ファイルを分割する数を決定
        num_parts = len(api_keys)  # APIキーの数に応じて分割数を決定

        transcribed_texts = [None] * num_parts  # インデックスに基づいて配置するリスト

        audio_parts = split_audio_file(audio_file_path, num_parts)

        with concurrent.futures.ThreadPoolExecutor() as executor:
            future_to_index = {executor.submit(transcribe_audio_with_key, part, api_keys[i]): i for i, part in enumerate(audio_parts)}
            failed_parts = []
            successful_api_keys = []  # 成功したAPIキーを記録するリスト
            for future in concurrent.futures.as_completed(future_to_index):
                index = future_to_index[future]
                part = audio_parts[index]
                result = future.result()
                if result:
                    transcribed_texts[index] = result
                    logging.info(f"{part}の処理が成功しました。")
                    if api_keys[index] not in successful_api_keys:
                        successful_api_keys.append(api_keys[index])  # 成功したAPIキーを記録
                else:
                    logging.error(f"{part}の処理が失敗しました。")
                    failed_parts.append((index, part))

        # 失敗したパートのリトライ
        if failed_parts:
            logging.info("失敗したファイルのリトライを1分後に開始します。")
            time.sleep(60)
            for index, part in failed_parts:
                result = transcribe_audio_with_key(part, api_keys[0])
                if result:
                    transcribed_texts[index] = result
                    logging.info(f"{part}のリトライが成功しました。")
                    if api_keys[0] not in successful_api_keys:
                        successful_api_keys.append(api_keys[0])  # リトライで成功したAPIキーを記録
                else:
                    logging.error(f"{part}のリトライが失敗しました。")

        # 分割されたファイルを削除
        for part in audio_parts:
            os.remove(part)
        logging.info(f"{audio_file_name}の分割されたファイルを削除しました。")

        # 文字起こし結果を結合（Noneを除外）
        combined_text = "\n".join(filter(None, transcribed_texts))
        # 余分な空白を取り除く
        cleaned_combined_text = " ".join(combined_text.split())
        logging.info(f"{audio_file_name}の文字起こしが完了しました。情報を抽出します。")

        # 文字起こし結果をWordファイルに保存
        try:
            output_directory = load_output_directory()
            word_output_file = os.path.join(output_directory, f"{os.path.splitext(audio_file_name)[0]}_文字起こし.docx")
            doc = Document()
            doc.add_paragraph(cleaned_combined_text)
            doc.save(word_output_file)
            logging.info(f"文字起こし結果がWordファイルに保存されました: {word_output_file}")
        except Exception as e:
            logging.error(f"文字起こし結果のWordファイル保存中にエラーが発生しました: {str(e)}")
            return False

        # 70秒のバッファを持たせる
        time.sleep(70)

        # 成功したAPIキーを使って情報抽出を試みる
        for api_key in successful_api_keys:
            try:
                extracted_info = extract_information(cleaned_combined_text, api_key)
                if extracted_info:
                    output_file = os.path.join(output_directory, f"{os.path.splitext(audio_file_name)[0]}_抽出結果.xlsx")
                    create_excel(extracted_info, output_file)
                    processed_files[audio_file_name] = output_file
                    break  # 成功したらループを抜ける
            except google.api_core.exceptions.ResourceExhausted:
                logging.error(f"{api_key}での情報抽出が失敗しました。次のAPIキーを試します。")
        else:
            logging.error(f"{audio_file_name}の情報抽出に失敗しました。")

        return True
    except Exception as e:
        logging.exception(f"{audio_file_path}の処理中にエラーが発生しました: {str(e)}")
        return False

def extract_info_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active
    data = {
        '会議名': sheet['B1'].value or '',
        '日時': convert_excel_date(sheet['B2'].value),
        '場所': sheet['B3'].value or '',
        '参加者': sheet['B4'].value or '',
        '欠席者': sheet['B5'].value or '',
    }
    for i in range(1, 11):  # 議題①から⑩まで
        data[f'議題{chr(0x2460 + i - 1)}'] = sheet[f'B{5+i*2-1}'].value or ''
        data[f'議題{chr(0x2460 + i - 1)}の要約'] = sheet[f'B{5+i*2}'].value or ''
    
    print("抽出されたデータ:")
    for key, value in data.items():
        print(f"{key}: {value}")
    
    return data

def convert_excel_date(value):
    if isinstance(value, (int, float)):
        return from_excel(value).strftime('%Y-%m-%d')
    try:
        # 文字列として日付を解析
        parsed_date = parser.parse(str(value))
        return parsed_date.strftime('%Y-%m-%d')
    except (ValueError, TypeError) as e:
        logging.error(f"日付の解析に失敗しました: {value} - {str(e)}")
        return value

def create_minutes_from_template(data, template_path):
    # 修正後
    template_path = os.path.join(get_current_dir(), 'テンプレート.docx')

    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        # 会議名、日時、場所、参加者、欠席者の置き換え
        for key, value in data.items():
            placeholder = f'「{key}」'
            if placeholder in paragraph.text:
                old_text = paragraph.text
                new_text = paragraph.text.replace(placeholder, str(value) if value is not None else '')
                paragraph.text = new_text
                print(f"置換: '{old_text}' -> '{new_text}'")

        # 議題と要約の置き換え
        for i in range(1, 11):
            topic_key = f'議題{chr(0x2460 + i - 1)}'
            topic_content = data.get(topic_key, '')
            summary_key = f'議題{chr(0x2460 + i - 1)}の要約'
            summary_content = data.get(summary_key, '')

            # 議題の名称を置き換え
            topic_placeholder = f'「{topic_key}」'
            if topic_placeholder in paragraph.text:
                old_text = paragraph.text
                new_text = paragraph.text.replace(topic_placeholder, topic_content)
                paragraph.text = new_text
                print(f"議題名置換: '{old_text}' -> '{new_text}'")

            # 要約の置き換え
            summary_placeholder = f'「{summary_key}」'
            if summary_placeholder in paragraph.text:
                old_text = paragraph.text
                new_text = paragraph.text.replace(summary_placeholder, summary_content)
                paragraph.text = new_text
                print(f"要約置換: '{old_text}' -> '{new_text}'")

    return doc

def create_minutes(xlsx_path, template_path, output_path):
    try:
        data = extract_info_from_xlsx(xlsx_path)
        doc = create_minutes_from_template(data, template_path)
        doc.save(output_path)
        print(f"議事録が作成されました: {output_path}")
        return True
    except Exception as e:
        logging.error(f"議事録の作成中にエラーが発生しました: {str(e)}")
        print(f"エラーが発生しました: {str(e)}")
        return False
    

# グローバル変数
selected_file = None
file_label = None
excel_file_label = None
uploading_label = None
elapsed_time_label = None
estimated_time_label = None
root = None
processing_done = False  # 処理が完了したかどうかを示すフラグ
start_time = None  # 処理開始時刻を保持
selected_file_name = ""  # 選択したファイル名を保持
estimated_time_text = ""  # 想定処理時間を保持

def show_main_menu():
    global root, file_label, excel_file_label, uploading_label, elapsed_time_label, estimated_time_label, selected_file, transcription_prompt, processing_done, start_time, selected_file_name, estimated_time_text
    selected_file = None

    # settings.jsonからプロンプトを再読み込み
    transcription_prompt = load_prompt_from_settings()
    logging.info(f"再読み込みしたプロンプト: {transcription_prompt}")

    for widget in root.winfo_children():
        widget.destroy()
 
    root.title("ファイル処理ツール")
    root.geometry("900x500")
    root.resizable(False, False)  # ウィンドウのサイズを固定
 
    # タイトルラベル
    title_label = tk.Label(root, text="⚡️爆速議事録", font=("Arial Black", 32, "bold"))
    title_label.pack(pady=20)
 
    # 設定ボタンを右上に配置
    settings_button = tk.Button(root, text="設定", command=show_settings, width=5, height=1)
    settings_button.place(x=800, y=20)  # 右上に配置

    # 使い方ボタンを設定ボタンの下に配置
    usage_button = tk.Button(root, text="使い方", command=show_usage, width=5, height=1)
    usage_button.place(x=800, y=60)  # 設定ボタンの下に配置

    # 音声ファイル処理フレーム
    audio_frame = tk.Frame(root, bd=2, relief="groove", width=350, height=400)
    audio_frame.pack_propagate(False)  # フレームのサイズを固定
    audio_frame.pack(side="left", padx=60, pady=20)  # 中央のスペースを縮める
 
    audio_label = tk.Label(audio_frame, text="音声ファイル処理", font=("Arial", 16, "bold"))
    audio_label.pack(pady=20)
 
    audio_button = tk.Button(audio_frame, text="音声ファイルを選択する", command=upload_audio_file)
    audio_button.pack(pady=10)

    # 選択したファイルを表示するラベル
    file_label = tk.Label(audio_frame, text=f"選択したファイル\n{selected_file_name}", wraplength=300, justify="center")
    file_label.pack(pady=10)

    # 音声ファイルを処理するボタン
    process_audio_button = tk.Button(audio_frame, text="音声ファイルを処理する", command=complete_audio_upload)
    process_audio_button.pack(pady=(20, 0))  # 初期位置を下げて固定
 
    # 想定処理時間を表示するラベル
    estimated_time_label = tk.Label(audio_frame, text=estimated_time_text, font=("Arial", 12))
    estimated_time_label.pack(pady=10)

    # 経過時間を表示するラベル
    uploading_label = tk.Label(audio_frame, text="", font=("Arial", 12))
    uploading_label.pack(pady=10)

    # 処理が進行中の場合、経過時間を更新
    if start_time and not processing_done:
        def update_elapsed_time():
            if not processing_done:
                elapsed_time = int(time.time() - start_time)
                minutes, seconds = divmod(elapsed_time, 60)
                if minutes > 0:
                    uploading_label.config(text=f"経過時間: {minutes}分{seconds}秒")
                else:
                    uploading_label.config(text=f"経過時間: {seconds}秒")
                root.after(1000, update_elapsed_time)  # 1秒ごとに更新

        update_elapsed_time()

    # Excelファイル処理フレーム
    excel_frame = tk.Frame(root, bd=2, relief="groove", width=350, height=400)
    excel_frame.pack_propagate(False)  # フレームのサイズを固定
    excel_frame.pack(side="right", padx=60, pady=20)  # 中央のスペースを縮める
 
    excel_label = tk.Label(excel_frame, text="Excelファイル処理", font=("Arial", 16, "bold"))
    excel_label.pack(pady=20)
 
    excel_button = tk.Button(excel_frame, text="Excelファイルを選択する", command=upload_xlsx_file)
    excel_button.pack(pady=10)
 
    excel_file_label = tk.Label(excel_frame, text="選択したファイル", wraplength=300, justify="center")
    excel_file_label.pack(pady=10)
 
    process_excel_button = tk.Button(excel_frame, text="Excelファイルを処理する", command=complete_xlsx_upload)
    process_excel_button.pack(pady=(20, 0))

def show_usage():
    for widget in root.winfo_children():
        widget.destroy()

    root.title("使い方")

    usage_label = tk.Label(root, text="使い方", font=("Arial", 16, "bold"))
    usage_label.pack(pady=(120, 0))  # 上に60ピクセルの余白を追加

    usage_text = "使い方は以下のWebページをご覧ください"
    usage_info = tk.Label(root, text=usage_text, justify="left")
    usage_info.pack(pady=10)

    # ホームページのリンク
    link = tk.Label(root, text="URLはこちら", fg="blue", cursor="hand2")
    link.pack(pady=10)
    link.bind("<Button-1>", lambda e: webbrowser.open("https://abiding-delivery-6d9.notion.site/1264d14a044c804f9dc7e41ce20a920f"))  # ここに実際のURLを入れてください

    usage_text = "爆速議事録をご利用いただきありがとうございます！"
    usage_info = tk.Label(root, text=usage_text, justify="left")
    usage_info.pack(pady=10)

    # 問い合わせのリンク
    link = tk.Label(root, text="お問い合わせ、バグの報告はこちらのアカウントまで", fg="blue", cursor="hand2")
    link.pack(pady=10)
    link.bind("<Button-1>", lambda e: webbrowser.open("https://x.com/petit_hiroto"))  # ここに実際のURLを入れてください

    # 戻るボタンを右上に配置
    back_button = tk.Button(root, text="戻る", command=show_main_menu, width=5, height=1)
    back_button.place(x=800, y=20)
 

def load_prompt_from_settings():
    """settings.jsonからプロンプトを読み込む関数"""
    settings_path = get_settings_path()
    logging.info(f"Settings path: {settings_path}")  # 追加: パスをログに出力
    if os.path.exists(settings_path):
        logging.info("settings.jsonが見つかりました。")  # 追加: ファイル存在確認
        with open(settings_path, 'r', encoding='utf-8') as f:
            try:
                settings = json.load(f)
                logging.info("settings.jsonを正常に読み込みました。")  # 追加: 読み込み成功
                return settings.get('transcription_prompt', '')  # デフォルト値を空文字に変更
            except json.JSONDecodeError as e:
                logging.error(f"JSONデコードエラー: {str(e)}")  # 追加: JSONデコードエラー
    else:
        logging.error("settings.jsonが見つかりません。")  # 追加: ファイルが見つからない場合
    return ''  # ファイルが存在しない場合も空文字を返す

def save_prompt_to_settings(prompt_text):
    """プロンプトをsettings.jsonに保存する関数"""
    ensure_settings_exist()  # フォルダとファイルの存在を確認
    settings_path = get_settings_path()
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        settings['transcription_prompt'] = prompt_text
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        logging.info("プロンプトがsettings.jsonに保存されました。")
        messagebox.showinfo("保存", "プロンプトが保存されました。")
    except Exception as e:
        logging.error(f"プロンプトの保存中にエラーが発生しました: {str(e)}")
        messagebox.showerror("エラー", "プロンプトの保存中にエラーが発生しました。")

def load_output_directory():
    """settings.jsonから出力先ディレクトリを読み込む関数"""
    settings_path = get_settings_path()
    if os.path.exists(settings_path):
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
            return settings.get('output_directory', '')  # デフォルト値を空文字に変更
    return ''  # ファイルが存在しない場合も空文字を返す

def save_output_directory_to_settings(directory):
    """出力先ディレクトリをsettings.jsonに保存する関数"""
    ensure_settings_exist()  # フォルダとファイルの存在を確認
    settings_path = get_settings_path()
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        settings['output_directory'] = directory
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        logging.info("出力先ディレクトリがsettings.jsonに保存されました。")
    except Exception as e:
        logging.error(f"出力先ディレクトリの保存中にエラーが発生しました: {str(e)}")

def show_settings():
    for widget in root.winfo_children():
        widget.destroy()
 
    root.title("設定")
    
    # 左半分のフレーム
    left_frame = tk.Frame(root, width=450, height=500)
    left_frame.pack_propagate(False)
    left_frame.pack(side="left", padx=30, pady=60)  # 中央に寄せるためにpadxを調整
 
    # 文字起こしプロンプトを変更する
    prompt_label = tk.Label(left_frame, text="文字起こしプロンプトを変更する", font=("Arial", 16, "bold"))
    prompt_label.grid(row=0, column=0, pady=10)

    # テキストボックスとスクロールバー
    prompt_textbox = tk.Text(left_frame, wrap="word", height=20, width=50)
    prompt_textbox.grid(row=1, column=0, sticky="nsew")

    # settings.jsonからプロンプトを読み込んで表示
    prompt_text = load_prompt_from_settings()
    if prompt_text:
        prompt_textbox.insert('1.0', prompt_text)

    scrollbar = tk.Scrollbar(left_frame, command=prompt_textbox.yview)
    scrollbar.grid(row=1, column=1, sticky="ns")
    prompt_textbox.config(yscrollcommand=scrollbar.set)
 
    # コピーアンドペーストを有効にする
    def enable_copy_paste(event):
        prompt_textbox.event_generate("<<Copy>>")
 
    prompt_textbox.bind("<Control-c>", enable_copy_paste)
    prompt_textbox.bind("<Control-v>", lambda e: prompt_textbox.event_generate("<<Paste>>"))
    prompt_textbox.bind("<Control-a>", lambda e: prompt_textbox.tag_add("sel", "1.0", "end-1c"))
 
    # 保存ボタンを左半分に追加
    save_prompt_button = tk.Button(left_frame, text="保存", command=lambda: save_prompt_to_settings(prompt_textbox.get('1.0', 'end-1c')))
    save_prompt_button.grid(row=2, column=0, pady=10)

    # 右半分のフレーム
    right_frame = tk.Frame(root, width=450, height=500)
    right_frame.pack_propagate(False)
    right_frame.pack(side="right", padx=30, pady=60)  # 中央に寄せるためにpadxを調整
 
    # 出力先ディレクトリを指定する
    directory_label = tk.Label(right_frame, text="出力先ディレクトリを指定する", font=("Arial", 16, "bold"))
    directory_label.grid(row=0, column=0, pady=10)

    # 現在指定されているディレクトリ
    current_dir = load_output_directory()
    current_dir_label = tk.Label(right_frame, text=f"指定されているディレクトリ\n:{current_dir}")
    current_dir_label.grid(row=1, column=0)

    # ディレクトリを指定するボタン
    def select_directory():
        directory = filedialog.askdirectory()
        if directory:
            # 選択されたディレクトリを表示
            current_dir_label.config(text=f"選択されたディレクトリ: {directory}")
            # ディレクトリパスをトリミングして保存
            save_output_directory_to_settings(directory.strip())

    directory_button = tk.Button(right_frame, text="ディレクトリを指定する", command=select_directory)
    directory_button.grid(row=3, column=0, pady=10)

    # 保存ボタン
    save_button = tk.Button(right_frame, text="保存", command=lambda: messagebox.showinfo("保存", "設定が保存されました。"))
    save_button.grid(row=4, column=0, pady=7)

    # Gemini APIキーを設定する
    api_key_label = tk.Label(right_frame, text="Gemini APIキーを設定する", font=("Arial", 16, "bold"))
    api_key_label.grid(row=5, column=0, pady=7)

    # テキストボックスとスクロールバー
    api_key_textbox = tk.Text(right_frame, wrap="word", height=10, width=50)
    api_key_textbox.grid(row=6, column=0, sticky="nsew")

    # settings.jsonからAPIキーを読み込んで表示
    api_keys_text = get_api_keys_text()
    if api_keys_text:
        api_key_textbox.insert('1.0', api_keys_text)

    api_key_scrollbar = tk.Scrollbar(right_frame, command=api_key_textbox.yview)
    api_key_scrollbar.grid(row=6, column=1, sticky="ns")
    api_key_textbox.config(yscrollcommand=api_key_scrollbar.set)

    # 保存ボタンを右半分に追加
    save_api_key_button = tk.Button(right_frame, text="保存", command=lambda: save_api_keys_to_settings(api_key_textbox.get('1.0', 'end-1c')))
    save_api_key_button.grid(row=7, column=0, pady=5)

    # 戻るボタンを右上に配置
    back_button = tk.Button(root, text="戻る", command=show_main_menu, width=5, height=1)
    back_button.place(x=800, y=20)

def upload_audio_file():
    global selected_file, selected_file_name, estimated_time_text
    selected_file = filedialog.askopenfilename(filetypes=[("Audio Files", "*.wav *.mp3 *.m4a")])
    if selected_file:
        selected_file_name = os.path.basename(selected_file)
        file_label.config(text=f"選択したファイル\n{selected_file_name}")
        
        # ファイルサイズを取得
        file_size_mb = os.path.getsize(selected_file) / (1024 * 1024)  # MBに変換
        
        # 想定処理時間を計算
        if file_size_mb <= 10:
            estimated_time = "1〜2分"
        elif file_size_mb <= 20:
            estimated_time = "2〜3分"
        else:
            estimated_time = "3〜5分"
        
        # 想定処理時間を表示
        estimated_time_text = f"想定処理時間：約{estimated_time}"
        estimated_time_label.config(text=estimated_time_text)

def complete_audio_upload():
    global start_time, processing_done
    if selected_file:
        start_time = time.time()  # 処理開始時刻を記録
        processing_done = False
        root.update_idletasks()
        processed_files = load_processed_files()
        threading.Thread(target=process_audio_file_async, args=(selected_file, processed_files, start_time)).start()
    else:
        messagebox.showwarning("警告", "ファイルが選択されていません。")

def upload_xlsx_file():
    global selected_file
    selected_file = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if selected_file:
        excel_file_label.config(text=f"選択したファイル\n{os.path.basename(selected_file)}")

def complete_xlsx_upload():
    if selected_file:
        root.update_idletasks()
        threading.Thread(target=process_xlsx_file_async, args=(selected_file,)).start()
    else:
        messagebox.showwarning("警告", "ファイルが選択されていません。")

def process_audio_file_async(audio_file, processed_files, start_time):
    global processing_done, selected_file, selected_file_name, estimated_time_text

    def update_elapsed_time():
        if not processing_done:
            elapsed_time = int(time.time() - start_time)
            minutes, seconds = divmod(elapsed_time, 60)
            if minutes > 0:
                uploading_label.config(text=f"経過時間: {minutes}分{seconds}秒")
            else:
                uploading_label.config(text=f"経過時間: {seconds}秒")
            root.after(1000, update_elapsed_time)  # 1秒ごとに更新

    threading.Thread(target=update_elapsed_time, daemon=True).start()

    # プロンプトが空でないか確認
    if not transcription_prompt:  # ここでグローバル変数を参照
        logging.error("プロンプトが空です。音声ファイルの処理を中止します。")
        processing_done = True  # 経過時間の更新を停止
        root.after(0, lambda: messagebox.showerror("エラー", "プロンプトが空です。処理を中止します。"))
        return  # 処理を中止

    try:
        logging.info(f"{audio_file}の処理を開始します。")
        success = process_audio_file(audio_file, processed_files)
        processing_done = True
        total_elapsed_time = int(time.time() - start_time)
        minutes, seconds = divmod(total_elapsed_time, 60)
        if minutes > 0:
            root.after(0, lambda: uploading_label.config(text=f"処理にかかった時間: {minutes}分{seconds}秒で処理が完了しました"))
        else:
            root.after(0, lambda: uploading_label.config(text=f"処理にかかった時間: {seconds}秒で処理が完了しました"))

        if success:
            # 処理が成功した場合、選択したファイル情報と想定処理時間をリセット
            root.after(0, lambda: reset_file_info())
            root.after(0, lambda: (messagebox.showinfo("完了", "ファイルのアップロードが完了しました。"), show_main_menu()))
        else:
            root.after(0, lambda: messagebox.showerror("エラー", "ファイルの処理中にエラーが発生しました。"))
    except Exception as e:
        logging.exception(f"音声ファイルの処理中にエラーが発生しました: {str(e)}")
        root.after(0, lambda: messagebox.showerror("エラー", "音声ファイルの処理中にエラーが発生しました。"))

def process_xlsx_file_async(xlsx_file):
    template_path = os.path.join(get_current_dir(), 'テンプレート.docx')  # dist直下から取得
    output_directory = load_output_directory()
    output_path = os.path.join(output_directory, f"{os.path.splitext(os.path.basename(xlsx_file))[0]}_議事録.docx")
    
    success = create_minutes(xlsx_file, template_path, output_path)
    if success:
        root.after(0, lambda: (messagebox.showinfo("完了", "議事録の作成が完了しました。"), show_main_menu()))
    else:
        messagebox.showerror("エラー", "ファイルの処理中にエラーが発生しました。")

def load_api_keys():
    """settings.jsonからAPIキーを読み込む関数"""
    settings_path = get_settings_path()  # 修正: get_current_dir() から get_settings_path() に変更
    if os.path.exists(settings_path):
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
            return [settings['gemini_api_keys'][f'GEMINI_API_KEY_{i}'] for i in range(1, 11)]
    logging.error("settings.jsonが見つからないか、APIキーが設定されていません。")
    return []

def get_api_keys_text():
    """APIキーをテキストボックスに表示するための文字列を生成する関数"""
    api_keys = load_api_keys()
    return "\n".join(api_keys)

def save_api_keys_to_settings(api_keys_text):
    """APIキーをsettings.jsonに保存する関数"""
    ensure_settings_exist()  # フォルダとファイルの存在を確認
    settings_path = get_settings_path()
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        api_keys = api_keys_text.strip().split('\n')
        settings['gemini_api_keys'] = {f'GEMINI_API_KEY_{i+1}': key for i, key in enumerate(api_keys)}
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        logging.info("APIキーがsettings.jsonに保存されました。")
        messagebox.showinfo("保存", "APIキーが保存されました。")
    except Exception as e:
        logging.error(f"APIキーの保存中にエラーが発生しました: {str(e)}")
        messagebox.showerror("エラー", "APIキーの保存中にエラーが発生しました。")

def main():
    global root, transcription_prompt  # グローバル変数を宣言
    try:
        logging.info("プロンプトをロード中...")  # 追加: ロード開始ログ
        transcription_prompt = load_prompt_from_settings()  # プロンプトをロード
        logging.info(f"取得したプロンプト: {transcription_prompt}")  # プロンプトの内容をログに出力
        logging.info("プロンプトのロードが完了しました。")  # 追加: ロード完了ログ
        root = tk.Tk()
        root.title("ファイル処理ツール")
        root.geometry("500x300")

        show_main_menu()

        root.mainloop()
    except Exception as e:
        logging.exception("アプリケーションの実行中にエラーが発生しました。")
        messagebox.showerror("エラー", f"アプリケーションの実行中にエラーが発生しました:\n{str(e)}")
        logging.error(f"アプリケーションの起動時にエラーが発生しました: {str(e)}")  # エラーログを追加

def get_settings_path():
    # ユーザーディレクトリのアプリケーションデータフォルダに保存
    return Path.home() / ".my_app" / "settings.json"

def load_settings():
    settings_path = get_settings_path()
    if settings_path.exists():
        with open(settings_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_settings(settings=None):
    try:
        settings_path = get_settings_path()
        settings_path.parent.mkdir(parents=True, exist_ok=True)  # フォルダが存在しない場合は作成

        # settings.jsonが存在しない場合、デフォルトの設定を作成
        if not settings_path.exists():
            settings = {
                'transcription_prompt': '',
                'output_directory': str(Path.home() / 'Documents'),
                'gemini_api_keys': {f'GEMINI_API_KEY_{i+1}': '' for i in range(10)}
            }

        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        print("Settings saved successfully.")  # ログ出力
    except Exception as e:
        print(f"Error saving settings: {e}")  # エラーログ

def ensure_settings_exist():
    settings_path = Path.home() / ".my_app" / "settings.json"
    
    # フォルダが存在しない場合は作成
    if not settings_path.parent.exists():
        settings_path.parent.mkdir(parents=True, exist_ok=True)
        print(f"フォルダを作成しました: {settings_path.parent}")
    
    # settings.jsonが存在しない場合は作成
    if not settings_path.exists():
        settings = {
            'transcription_prompt': '',
            'output_directory': '',
            'gemini_api_keys': {f'GEMINI_API_KEY_{i+1}': '' for i in range(10)}
        }
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        print(f"settings.jsonを作成しました: {settings_path}")
    else:
        print(f"settings.jsonは既に存在します: {settings_path}")

# 確認と作成を実行
ensure_settings_exist()

def reset_file_info():
    global selected_file, selected_file_name, estimated_time_text
    selected_file = None
    selected_file_name = ""
    estimated_time_text = ""
    file_label.config(text="選択したファイル\n")
    estimated_time_label.config(text="")

if __name__ == "__main__":
    main()

